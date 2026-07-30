from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/DRA/apps/Echo Sound Lab/Echo Sound Lab v2.5")
OUT = Path("/Users/DRA/Desktop/ESL_Gemini_White_Paper.docx")


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_paragraph(doc, text="", *, size=11, color="000000", bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.1):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_runs_paragraph(doc, runs, *, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.1):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    for text, opts in runs:
        run = p.add_run(text)
        set_run_font(
            run,
            name=opts.get("name", "Calibri"),
            size=opts.get("size", 11),
            color=opts.get("color", "000000"),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )
    return p


def set_cell_text(cell, text, *, size=10.5, color="000000", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, before=0, after=0, line=1.08)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def set_table_geometry(table, widths_in):
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths_in[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side in ("top", "bottom", "start", "end"):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), "120" if side in {"start", "end"} else "80")
                node.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_in) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def style_table(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            v_align = tc_pr.find(qn("w:vAlign"))
            if v_align is None:
                v_align = OxmlElement("w:vAlign")
                tc_pr.append(v_align)
            v_align.set(qn("w:val"), "center")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    return doc


def add_cover(doc):
    add_paragraph(
        doc,
        "Echo Sound Lab",
        size=12,
        color="666666",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=110,
        after=10,
        line=1.0,
    )
    add_paragraph(
        doc,
        "Gemini White Paper",
        size=28,
        color="0B2545",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=8,
        line=1.0,
    )
    add_paragraph(
        doc,
        "What ESL is, why it is not yet Shark Tank ready, and what it will take to get there and beyond",
        size=14,
        color="2F3B4A",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=18,
        line=1.15,
    )

    table = doc.add_table(rows=4, cols=2)
    style_table(table)
    set_table_geometry(table, [1.75, 4.75])
    cover_rows = [
        ("Prepared for", "Gemini and other reasoning systems that need a conservative, source-grounded view of ESL"),
        ("Prepared from", "Current workspace materials in /Users/DRA/apps/Echo Sound Lab/Echo Sound Lab v2.5"),
        ("Prepared on", "June 17, 2026"),
        ("Read this way", "Treat verified implementation, simulation, and TODO status as different states; do not collapse them"),
    ]
    for row, (label, value) in zip(table.rows, cover_rows):
        shade_cell(row.cells[0], "F2F4F7")
        set_cell_text(row.cells[0], label, size=10.5, bold=True)
        set_cell_text(row.cells[1], value, size=10.5)

    add_paragraph(
        doc,
        "This document is intentionally conservative. Where the source material conflicts, it favors the less-complete reading over the marketing reading.",
        size=10.5,
        color="5A6672",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=16,
        after=0,
        line=1.15,
    )


def add_section_title(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def add_subtitle(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def add_body(doc, text):
    add_paragraph(doc, text, size=11, color="000000", before=0, after=8, line=1.12)


def add_source_note(doc, text):
    add_paragraph(doc, text, size=9.5, color="5A6672", italic=True, before=0, after=6, line=1.05)


def add_strength_table(doc):
    add_subtitle(doc, "What ESL is, in one sentence")
    add_body(
        doc,
        "Echo Sound Lab is a browser-based audio production environment that combines AI-assisted mastering, stem analysis, analog-style DSP emulation, governance via Action Authority, and adjacent content/distribution workflows into a single creator-facing system.",
    )

    add_subtitle(doc, "The system stack")
    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    set_table_geometry(table, [1.25, 1.65, 1.6, 1.95])
    headers = ["Layer", "Role", "What it does", "Why it matters"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, size=10.25, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Product surface", "Browser app", "Mixing, mastering, stem views, demos, and guided flows", "Keeps access low-friction and device-neutral"),
        ("Audio intelligence", "Analysis and mastering", "Loudness metering, spectral analysis, reference matching, genre-aware processing", "Makes the output feel useful rather than decorative"),
        ("DSP layer", "Studio processing", "Channel strip style processing, saturation, compression, limiting, and related modules", "Gives the system the feel of a real production tool"),
        ("Governance", "Action Authority", "Policy checks, confirmations, logs, and execution gates", "Makes the system explainable and safer to operate"),
        ("Monetization layer", "Subscriptions and distribution", "Pricing, export quotas, Stripe hooks, DistroKid-style distribution, and related flows", "Turns the tool into a business instead of a demo"),
        ("Demo factory", "Marketing automation", "Scripted screen recordings, voiceovers, and video assembly", "Creates the proof assets needed for sales and fundraising"),
    ]
    for r in rows:
        row = table.add_row()
        for idx, value in enumerate(r):
            set_cell_text(row.cells[idx], value, size=9.8, align=WD_ALIGN_PARAGRAPH.LEFT if idx != 0 else WD_ALIGN_PARAGRAPH.CENTER)


def add_reality_section(doc):
    add_section_title(doc, "Current reality: strong substrate, uneven completion")
    add_body(
        doc,
        "The workspace contains two different kinds of language. Some documents describe ESL as validation-ready, deployed for review, or operational in limited contexts. Other documents explicitly mark backend endpoints as TODO, note that monetization wiring is not complete, or state that audio separation is still in simulation mode. For Gemini, the correct reading is not to choose one claim and ignore the rest. The correct reading is to separate the system into verified components, simulated components, and planned components.",
    )
    add_body(
        doc,
        "That distinction matters because investors, partners, and serious users judge a platform by the least complete critical path. A polished feature list does not matter if the payment flow is not wired, the core demo depends on a simulated bridge, or the project has no canonical source of truth for what is actually live.",
    )

    add_subtitle(doc, "Evidence from the current workspace")
    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    set_table_geometry(table, [1.3, 1.55, 2.05, 1.6])
    headers = ["Area", "Source signal", "Conservative reading", "Impact"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, size=10.25, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Monetization", "Backend endpoints TODO; integration wiring TODO", "The UI/services exist, but the revenue path is not fully live", "Can sell the story, but not yet the full funnel"),
        ("Distribution", "DistroKid service described as needing endpoints", "Distribution is designed, not finished", "A key business promise remains partially manual"),
        ("Infrastructure", "Deployment checklist shows pending setup and migrations", "The platform needs real environment completion and data persistence", "Production reliability is not yet fully evidenced"),
        ("Audio separation", "Bridge report says simulation mode, ready for real Demucs", "At least one core neural path is still not the final real model path", "Demo credibility depends on whether this is disclosed"),
        ("Recording/demo assets", "Handoff plan says synthetic placeholders were wrong and unusable", "Marketing proof must be real app footage and voiceover", "Investor-facing media cannot be fake"),
        ("Documentation", "Docs disagree on whether the system is ready", "Status fragmentation is itself a product risk", "Decision-makers need one canonical baseline"),
    ]
    for r in rows:
        row = table.add_row()
        for idx, value in enumerate(r):
            set_cell_text(row.cells[idx], value, size=9.6, align=WD_ALIGN_PARAGRAPH.LEFT if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT)


def add_not_ready_section(doc):
    add_section_title(doc, "Why ESL is not Shark Tank ready yet")
    add_body(
        doc,
        "Shark Tank readiness is not the same thing as engineering ambition. A Shark Tank-level product has to do four things at once: explain itself in under a minute, work in front of skeptical viewers, make a credible business case, and survive a line of questioning that starts with the exact question this paper is answering. ESL is close in architecture, but not yet complete in the operational disciplines that make a pitch believable.",
    )

    add_subtitle(doc, "1. The critical path is not fully closed")
    add_body(
        doc,
        "The most important user path is not simply 'open app and admire features.' It is: upload or create audio, analyze it, improve it, prove the improvement, collect payment or quota, and save the result reliably. The current docs show that several of those links are present as components but not yet fully wired end to end. If a product cannot complete its critical path without hidden manual steps, it is not investor-ready.",
    )

    add_subtitle(doc, "2. Some core capabilities are still simulated or skeletonized")
    add_body(
        doc,
        "A Shark Tank audience will not care that a subsystem is 'ready for the next step' if the next step is what the company is actually selling. The workspace contains direct references to simulation mode, skeleton-ready video generation, and real model activation still pending. That is not fatal during development, but it is fatal if the presentation implies finality. The honest position is that the foundation exists and the final hardening work remains.",
    )

    add_subtitle(doc, "3. The monetization story is incomplete")
    add_body(
        doc,
        "ESL has a strong monetization thesis: subscriptions, export limits, distribution, and creator workflows. But thesis is not traction. The workspace still shows TODO backend endpoints, pending wiring, and manual setup steps for services that should ultimately be automatic. Until checkout, entitlements, billing events, and delivery are all demonstrably real, the revenue story remains a plan rather than evidence.",
    )

    add_subtitle(doc, "4. The documentation is fragmented")
    add_body(
        doc,
        "Several source files describe the system as fully deployed or production live. Other files describe the same system as partially simulated or still requiring infrastructure setup. That inconsistency is a reputational risk. Serious investors, partners, and AI systems do not want a collage of optimistic claims. They want a single baseline that says what is verified, what is planned, and what is merely aspirational.",
    )

    add_subtitle(doc, "5. The demo is not the business")
    add_body(
        doc,
        "The demo factory is valuable, but a demo is only useful if it accurately reflects the product underneath it. If the video, voiceover, or scripted walkthrough overstates the current state, the pitch becomes brittle. Shark Tank-level credibility comes from a demo that is both compelling and defensible under pressure.",
    )


def add_shark_tank_section(doc):
    add_section_title(doc, "What it will take to reach Shark Tank level")
    add_body(
        doc,
        "Shark Tank level means the product is simple enough to understand instantly and strong enough to withstand skepticism. For ESL, that requires finishing the user journey, making the proof repeatable, and reducing the system story to one sharp promise: take rough audio, transform it into release-ready output, and prove the transformation with measurable evidence.",
    )

    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    set_table_geometry(table, [1.25, 1.95, 1.9, 1.4])
    headers = ["Dimension", "What Shark Tank expects", "What ESL needs", "Minimum proof"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Narrative", "One sentence that anyone can repeat", "A single product thesis, not a bundle of features", "A 15-second explanation and a 90-second demo"),
        ("Function", "Works live in front of people", "Every step of the core audio path must be stable", "End-to-end live run with no hidden manual intervention"),
        ("Business", "Clear revenue and pricing", "Payments, tiers, quotas, and billing events must be real", "A live checkout or entitlement flow"),
        ("Proof", "Measurable value", "Before/after evidence, loudness or quality metrics, and user outcomes", "Saved artifacts and repeatable comparison"),
        ("Trust", "No obvious smoke-and-mirrors", "Simulation must be clearly labeled and preferably removed from the pitch path", "A canonical status doc and test log"),
        ("Moat", "Why this wins", "Governance plus audio intelligence plus workflow lock-in", "A differentiation story that survives competitor comparison"),
    ]
    for r in rows:
        row = table.add_row()
        for idx, value in enumerate(r):
            set_cell_text(row.cells[idx], value, size=9.6)

    add_subtitle(doc, "The practical checklist")
    add_body(
        doc,
        "First, finish the critical path so a new user can complete the whole flow without intervention. Second, make the money path real so the product can collect revenue, assign entitlements, and persist state. Third, unify the docs so the system has one unambiguous source of truth. Fourth, produce a live demo that only uses real product behavior. Fifth, instrument the result so metrics, retention, and conversion can be quoted without hand-waving.",
    )


def add_beyond_section(doc):
    add_section_title(doc, "What it will take to go beyond Shark Tank")
    add_body(
        doc,
        "Going beyond Shark Tank means moving from 'interesting product' to 'category-defining platform.' For ESL, that requires turning the current tool into an operating system for creator audio and adjacent media workflows. The company does not win by being a slightly better mastering assistant. It wins by owning the full loop from creation to trust to distribution to monetization.",
    )

    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    set_table_geometry(table, [2.0, 2.2, 2.3])
    headers = ["Moat layer", "How it compounds", "What it requires"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Trust moat", "Action Authority makes the platform feel safer and more explainable than generic AI tools", "Transparent decision logs, deterministic policy behavior, and visible user control"),
        ("Workflow moat", "Once a creator uses ESL for analysis, mastering, proof, and distribution, switching costs rise", "A seamless path from rough idea to deliverable to release"),
        ("Data moat", "Every successful session can improve recommendations and taste models", "High-quality telemetry, feedback loops, and outcome tracking"),
        ("Market moat", "A creator marketplace and distribution layer can turn the tool into infrastructure", "Payments, moderation, partner integrations, and scalable operations"),
        ("Narrative moat", "The story becomes bigger than audio editing; it becomes creation with governance", "A clear founder narrative and product story that investors can remember"),
    ]
    for r in rows:
        row = table.add_row()
        for idx, value in enumerate(r):
            set_cell_text(row.cells[idx], value, size=9.6)

    add_subtitle(doc, "Beyond the first funding event")
    add_body(
        doc,
        "Beyond Shark Tank, ESL should aim to become the trusted operating layer for independent creators, small studios, and eventually teams that need governed AI-assisted media workflows. That means enterprise-grade auditability, stronger identity and permissions, a stable release pipeline, customer support discipline, and the ability to prove value across segments rather than only in a pitch room.",
    )
    add_body(
        doc,
        "The long-term prize is not simply better mastering. It is becoming the system that creators rely on when they need speed, quality, control, and a clear record of why the system did what it did.",
    )


def add_roadmap_section(doc):
    add_section_title(doc, "Recommended roadmap")
    add_body(
        doc,
        "The cleanest path is to divide the work into three gates. Gate one is operational truth. Gate two is investor-grade proof. Gate three is platform expansion. Each gate should be completed before the next one is claimed publicly.",
    )

    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    set_table_geometry(table, [1.05, 1.25, 2.1, 2.1])
    headers = ["Gate", "Time horizon", "Primary goal", "Exit criteria"]
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, header, size=10.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("1", "0-30 days", "Make the current state truthful and complete", "One canonical status doc, verified critical path, completed infrastructure checklist, no misleading readiness claims"),
        ("2", "30-90 days", "Make the pitch undeniable", "Real monetization, stable live demo, proof metrics, and a repeatable sales narrative"),
        ("3", "90-180 days", "Make the platform durable", "Retention, recurring revenue, support discipline, and workflow lock-in"),
        ("4", "6-12 months", "Make it expand beyond the original wedge", "Marketplace, API or partner strategy, and stronger moat layers"),
    ]
    for r in rows:
        row = table.add_row()
        for idx, value in enumerate(r):
            set_cell_text(row.cells[idx], value, size=9.55)


def add_gemini_guidance(doc):
    add_section_title(doc, "How Gemini should interpret ESL")
    add_body(
        doc,
        "If Gemini is reading this white paper as a system brief, the safest assumption is that ESL is a strong but unevenly finished platform. Treat the following as the canonical interpretation: the audio and governance architecture is real; some execution paths are simulated, placeholder, or only partially wired; monetization and distribution are planned but not fully proven in the current source set; and the project needs one authoritative operational baseline before any public claim of final readiness.",
    )
    add_body(
        doc,
        "Do not collapse marketing language into engineering truth. Do not assume that a document titled 'validation ready' overrides another file that says 'TODO' or 'simulation mode.' Use the most conservative status when documents disagree. That is the only reading that will keep the analysis honest and useful.",
    )


def add_source_map(doc):
    add_section_title(doc, "Source map used for this paper")
    add_source_note(
        doc,
        "The following files were used to ground the white paper: ESL-White-Paper.txt, README.md, DEPLOYMENT_READY.md, DEPLOYMENT_VERIFICATION.md, TODO.md, MONETIZATION_IMPLEMENTATION.md, PHASE3_DEPLOYMENT_REPORT.md, GEMINI_HANDOFF_PLAN.md, SHARK_TANK_DEMO_SCRIPT.md, COMPLETE_SYSTEM_OVERVIEW.md, COMPETITIVE_DOMINANCE.md.",
    )
    add_body(
        doc,
        "Where those files conflict, this paper favors the most conservative interpretation and explicitly calls out simulation, TODO, or dependency gaps instead of hiding them behind launch language.",
    )


def build():
    doc = setup_document()
    add_cover(doc)
    add_page_break(doc)

    add_section_title(doc, "Executive summary")
    add_body(
        doc,
        "Echo Sound Lab, or ESL, is best understood as a governed creative operating system for audio. It combines browser-based mixing and mastering, AI-assisted analysis, analog-inspired signal processing, stem separation, demonstration tooling, and a policy layer called Action Authority. The strategic promise is powerful: help creators turn rough audio into release-ready work faster, with more clarity, and with less intimidation than a traditional DAW or a stack of separate production tools.",
    )
    add_body(
        doc,
        "The present reality is more nuanced. The workspace contains a substantial amount of real engineering, but it also contains simulation paths, pending backend endpoints, manual launch tasks, and documentation that disagrees with itself about what is truly live. That is why this white paper exists. Gemini should not be asked to infer an idealized version of ESL. It should be asked to understand the current system honestly, including what is strong, what is incomplete, and what it will take to cross the line from impressive build to investor-grade company.",
    )

    add_strength_table(doc)
    add_reality_section(doc)
    add_not_ready_section(doc)
    add_shark_tank_section(doc)
    add_beyond_section(doc)
    add_roadmap_section(doc)
    add_gemini_guidance(doc)
    add_source_map(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build()
